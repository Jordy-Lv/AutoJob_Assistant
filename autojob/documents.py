from __future__ import annotations

import html
import unicodedata
from pathlib import Path

from .config import OUTPUT_DIR, ensure_directories
from .models import AnalysisResult, GeneratedDocument, JobOffer, UserProfile


def slugify(value: str, fallback: str = "documento") -> str:
    normalized = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode()
    slug = "".join(char.lower() if char.isalnum() else "-" for char in normalized)
    slug = "-".join(part for part in slug.split("-") if part)
    return slug[:90] or fallback


def _require_docx():
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Instala python-docx para generar archivos DOCX.") from exc
    return Document


def _require_reportlab():
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError as exc:
        raise RuntimeError("Instala reportlab para generar archivos PDF.") from exc
    return {
        "colors": colors,
        "letter": letter,
        "styles": getSampleStyleSheet,
        "ListFlowable": ListFlowable,
        "ListItem": ListItem,
        "Paragraph": Paragraph,
        "SimpleDocTemplate": SimpleDocTemplate,
        "Spacer": Spacer,
    }


def _non_empty_lines(value: str) -> list[str]:
    return [line.strip() for line in value.splitlines() if line.strip()]


def _add_bullets_docx(document, title: str, items: list[str]) -> None:
    if not items:
        return
    document.add_heading(title, level=2)
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _cv_path(job: JobOffer, suffix: str) -> Path:
    base = slugify(f"{job.company}-{job.title}-cv", fallback="cv")
    return OUTPUT_DIR / f"{base}.{suffix}"


def _letter_path(job: JobOffer, suffix: str) -> Path:
    base = slugify(f"{job.company}-{job.title}-carta", fallback="carta")
    return OUTPUT_DIR / f"{base}.{suffix}"


def generate_cv_docx(
    profile: UserProfile,
    job: JobOffer,
    analysis: AnalysisResult,
) -> GeneratedDocument:
    ensure_directories()
    Document = _require_docx()
    document = Document()
    document.add_heading(profile.full_name or "Candidato", level=0)
    document.add_paragraph(profile.links)
    document.add_paragraph(f"Objetivo: {job.title} en {job.company or 'la empresa'}")

    if profile.summary:
        document.add_heading("Perfil", level=2)
        document.add_paragraph(profile.summary)

    document.add_heading("Ajuste para la oferta", level=2)
    document.add_paragraph(f"Compatibilidad estimada: {analysis.score}%")
    document.add_paragraph(analysis.recommendation)
    _add_bullets_docx(document, "Fortalezas relevantes", analysis.reasons)

    if analysis.matched_skills:
        _add_bullets_docx(document, "Habilidades destacadas", analysis.matched_skills)
    elif profile.skills:
        _add_bullets_docx(document, "Habilidades", profile.skills_list())

    _add_bullets_docx(document, "Experiencia", _non_empty_lines(profile.experience))
    _add_bullets_docx(document, "Proyectos", _non_empty_lines(profile.projects))
    _add_bullets_docx(document, "Educacion", _non_empty_lines(profile.education))

    if analysis.gaps:
        _add_bullets_docx(
            document,
            "Temas a reforzar antes de entrevista",
            analysis.gaps,
        )

    path = _cv_path(job, "docx")
    document.save(path)
    return GeneratedDocument("CV DOCX", str(path))


def generate_letter_docx(
    profile: UserProfile,
    job: JobOffer,
    analysis: AnalysisResult,
) -> GeneratedDocument:
    ensure_directories()
    Document = _require_docx()
    document = Document()
    document.add_heading("Carta de presentacion", level=0)
    document.add_paragraph(f"Para: {job.company or 'Equipo de seleccion'}")
    document.add_paragraph(f"Cargo: {job.title}")
    document.add_paragraph("")

    name = profile.full_name or "Jordy Pardo"
    intro_role = profile.target_role or "desarrollador de software"
    company = job.company or "su equipo"
    document.add_paragraph(
        f"Hola {company},"
    )
    document.add_paragraph(
        f"Me interesa postularme al rol de {job.title}. Soy {intro_role} y "
        "me motiva construir soluciones practicas, mantenibles y orientadas a resultados."
    )
    if analysis.matched_skills:
        document.add_paragraph(
            "Al revisar la oferta, vi una buena alineacion con mi experiencia en "
            + ", ".join(analysis.matched_skills[:6])
            + "."
        )
    if profile.summary:
        document.add_paragraph(profile.summary)

    document.add_paragraph(
        "Me gustaria conversar sobre como puedo aportar al equipo y al producto. "
        "Quedo atento a la posibilidad de una entrevista."
    )
    document.add_paragraph(f"Saludos,\n{name}")

    path = _letter_path(job, "docx")
    document.save(path)
    return GeneratedDocument("Carta DOCX", str(path))


def _paragraph(text: str, styles, name: str = "BodyText"):
    Paragraph = _require_reportlab()["Paragraph"]
    return Paragraph(html.escape(text), styles[name])


def _pdf_bullets(title: str, items: list[str], kit, story, styles) -> None:
    if not items:
        return
    Paragraph = kit["Paragraph"]
    Spacer = kit["Spacer"]
    ListFlowable = kit["ListFlowable"]
    ListItem = kit["ListItem"]
    story.append(Paragraph(html.escape(title), styles["Heading2"]))
    story.append(
        ListFlowable(
            [
                ListItem(Paragraph(html.escape(item), styles["BodyText"]))
                for item in items
            ],
            bulletType="bullet",
        )
    )
    story.append(Spacer(1, 8))


def generate_cv_pdf(
    profile: UserProfile,
    job: JobOffer,
    analysis: AnalysisResult,
) -> GeneratedDocument:
    ensure_directories()
    kit = _require_reportlab()
    styles = kit["styles"]()
    Paragraph = kit["Paragraph"]
    Spacer = kit["Spacer"]
    SimpleDocTemplate = kit["SimpleDocTemplate"]

    path = _cv_path(job, "pdf")
    doc = SimpleDocTemplate(str(path), pagesize=kit["letter"])
    story = [
        Paragraph(html.escape(profile.full_name or "Candidato"), styles["Title"]),
        Paragraph(html.escape(profile.links), styles["BodyText"]),
        Spacer(1, 8),
        Paragraph(
            html.escape(f"Objetivo: {job.title} en {job.company or 'la empresa'}"),
            styles["BodyText"],
        ),
        Spacer(1, 12),
    ]
    if profile.summary:
        story.append(Paragraph("Perfil", styles["Heading2"]))
        story.append(Paragraph(html.escape(profile.summary), styles["BodyText"]))
        story.append(Spacer(1, 8))

    story.append(Paragraph("Ajuste para la oferta", styles["Heading2"]))
    story.append(Paragraph(html.escape(f"Compatibilidad estimada: {analysis.score}%"), styles["BodyText"]))
    story.append(Paragraph(html.escape(analysis.recommendation), styles["BodyText"]))
    story.append(Spacer(1, 8))

    _pdf_bullets("Fortalezas relevantes", analysis.reasons, kit, story, styles)
    _pdf_bullets("Habilidades destacadas", analysis.matched_skills or profile.skills_list(), kit, story, styles)
    _pdf_bullets("Experiencia", _non_empty_lines(profile.experience), kit, story, styles)
    _pdf_bullets("Proyectos", _non_empty_lines(profile.projects), kit, story, styles)
    _pdf_bullets("Educacion", _non_empty_lines(profile.education), kit, story, styles)
    _pdf_bullets("Temas a reforzar antes de entrevista", analysis.gaps, kit, story, styles)
    doc.build(story)
    return GeneratedDocument("CV PDF", str(path))


def generate_letter_pdf(
    profile: UserProfile,
    job: JobOffer,
    analysis: AnalysisResult,
) -> GeneratedDocument:
    ensure_directories()
    kit = _require_reportlab()
    styles = kit["styles"]()
    Paragraph = kit["Paragraph"]
    Spacer = kit["Spacer"]
    SimpleDocTemplate = kit["SimpleDocTemplate"]

    path = _letter_path(job, "pdf")
    name = profile.full_name or "Jordy Pardo"
    intro_role = profile.target_role or "desarrollador de software"
    company = job.company or "su equipo"
    matched = ", ".join(analysis.matched_skills[:6])
    story = [
        Paragraph("Carta de presentacion", styles["Title"]),
        Paragraph(html.escape(f"Para: {company}"), styles["BodyText"]),
        Paragraph(html.escape(f"Cargo: {job.title}"), styles["BodyText"]),
        Spacer(1, 12),
        Paragraph(html.escape(f"Hola {company},"), styles["BodyText"]),
        Paragraph(
            html.escape(
                f"Me interesa postularme al rol de {job.title}. Soy {intro_role} "
                "y me motiva construir soluciones practicas, mantenibles y orientadas a resultados."
            ),
            styles["BodyText"],
        ),
    ]
    if matched:
        story.append(
            Paragraph(
                html.escape(
                    "Al revisar la oferta, vi una buena alineacion con mi experiencia en "
                    f"{matched}."
                ),
                styles["BodyText"],
            )
        )
    if profile.summary:
        story.append(Paragraph(html.escape(profile.summary), styles["BodyText"]))
    story.extend(
        [
            Paragraph(
                html.escape(
                    "Me gustaria conversar sobre como puedo aportar al equipo y al producto. "
                    "Quedo atento a la posibilidad de una entrevista."
                ),
                styles["BodyText"],
            ),
            Spacer(1, 12),
            Paragraph(html.escape(f"Saludos,\n{name}"), styles["BodyText"]),
        ]
    )
    doc = SimpleDocTemplate(str(path), pagesize=kit["letter"])
    doc.build(story)
    return GeneratedDocument("Carta PDF", str(path))


def generate_document_package(
    profile: UserProfile,
    job: JobOffer,
    analysis: AnalysisResult,
) -> list[GeneratedDocument]:
    return [
        generate_cv_docx(profile, job, analysis),
        generate_cv_pdf(profile, job, analysis),
        generate_letter_docx(profile, job, analysis),
        generate_letter_pdf(profile, job, analysis),
    ]
